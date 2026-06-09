#!/usr/bin/env python3
"""Generate Word doc for Green Tea Group Follow-up Research Report."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
def set_font(run, name_en="Times New Roman", name_zh="宋体", size=11, bold=False, color=None):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_zh)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(10,10,11)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level==1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    size = {1:16, 2:13, 3:11}.get(level, 11)
    bold = level <= 2
    set_font(run, size=size, bold=bold, color=color)
    return p

def body(text, indent=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(139, 108, 66))
    return p

def callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    # left border via shading
    run = p.add_run(text)
    set_font(run, size=10, color=(80,60,30))
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "8B6C42")
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F8F4EC")
    pPr.append(shd)
    return p

def confidence_badge(text, level):
    colors = {"high":(42,92,63), "mid":(139,100,40), "low":(139,46,46)}
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"[{text}]")
    set_font(run, size=9, bold=True, color=colors.get(level,(80,80,80)))
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(241,239,234))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0A0A0B")
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        tr = table.rows[ri+1]
        fill = "F8F4EC" if ri%2==0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=9.5)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A84C")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MERLIN INTELLIGENCE · MODE B")
set_font(run, size=8, color=(139,108,66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("绿茶集团延伸问题深度研究")
set_font(run, name_en="Times New Roman", name_zh="宋体", size=22, bold=True, color=(10,10,11))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Green Tea Group (06831.HK) · Follow-up Research Report")
set_font(run, size=10, color=(80,60,30))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("研究日期：2026年6月9日 ｜ 置信分布：高置信 1 / 中置信 5 / 低置信 2")
set_font(run, size=9, color=(100,100,100))

hr()

# ── Tags row ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tags = ["餐饮连锁", "品牌营销", "新品上市", "培训体系", "组织协同", "港股 06831.HK"]
for tag in tags:
    run = p.add_run(f" {tag} ")
    set_font(run, size=8.5, bold=True, color=(241,239,234))
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "0A0A0B")
    rPr.append(shd)
    p.add_run("  ")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Q1: 年度品牌日历审批
# ══════════════════════════════════════════════════════════════════════════════
heading("Q1 · 年度品牌日历审批层级与部门协同机制", 1)

confidence_badge("中置信 · MID CONFIDENCE", "mid")
body("绿茶集团的年度品牌日历审批流程，基于其已披露的三级组织架构（集团总部→区域→门店）和上市公司治理要求推断，具体会议名称与参会岗位需一手验证。")
doc.add_paragraph()

# 4-tier table
subheading("审批层级结构")
add_table(
    ["层级", "角色/岗位", "职责", "审批事项"],
    [
        ["L1 执行层", "品牌运营专员 / 社媒编辑", "日历初稿编制、内容策划、供应商沟通", "节日节点、促销方案草案"],
        ["L2 职能层", "品牌总监 / 营销总监 / 产品总监", "多职能对齐、资源协调、预算确认", "全年重要节点、跨部门资源分配"],
        ["L3 管理层", "CMO / VP品牌（于丽影副总裁）/ CEO（王勤松）", "战略一致性审核、重大活动最终拍板", "品牌主题、全国性大促、与产品线联动的节点"],
        ["L4 董事会层", "董事会 / 审计委员会（上市后要求）", "影响财务披露或品牌重大转向的决策", "超过一定金额的营销预算、品牌重大重塑"],
    ],
    [3.5, 4.5, 5.0, 4.5]
)

callout("⚠ 低置信提示：L4 董事会直接审批的触发条件（金额门槛、品牌转型定义）是推断，需向CFO或IR确认。港股上市公司（2025年5月上市）治理要求加重了对L3-L4连接机制的关注。")

subheading("跨部门协同机制")
add_table(
    ["部门", "在品牌日历中的角色", "与品牌部的接口"],
    [
        ["品牌/营销部", "主导方：编制日历、统筹执行", "内部发起、跨部门对齐枢纽"],
        ["产品研发部", "确认新品上市时间窗口，确保上线节点与供应链匹配", "提供产品roadmap输入，确认15-20天执行窗口"],
        ["运营/门店部", "确认门店执行能力，给出旺季/淡季门店侧反馈", "接收培训材料、SOP，参与节假日排班规划"],
        ["供应链/采购部", "与促销食材、活动物料需求对齐", "提前60-90天告知大促预估备货量"],
        ["数字化/私域部", "确认私域运营节奏与品牌日历同步", "输出私域推广档期计划，匹配公域投放时间"],
    ],
    [3.5, 5.0, 5.0]
)

hr()

# ══════════════════════════════════════════════════════════════════════════════
# Q2: 新品上市15-20天压缩
# ══════════════════════════════════════════════════════════════════════════════
heading("Q2 · 新品上市周期压缩至15-20天的并行执行机制", 1)

confidence_badge("中置信 · MID CONFIDENCE", "mid")
body("2024年绿茶集团新品203道（约4道/周），印证了高频并行上新能力的存在。以下执行框架基于行业标准和公开披露的数字化能力推断，15-20天数字来源需向产品/营销团队确认。")
doc.add_paragraph()

callout("【概念澄清】15-20天指的是「菜品内容制作与门店上线执行」的最后一公里窗口，不含前端研发、测试、供应链打通（那部分通常6-12周）。这是执行交付阶段的压缩，不是全链路R&D的压缩。")

subheading("Day-by-Day 并行执行流程")
add_table(
    ["阶段", "时间节点", "执行动作", "责任方"],
    [
        ["产品定型 + 拍摄启动", "Day 0-2", "菜品定型确认（主厨/研发）→ 通知拍摄团队备档期 → 确认摄影师/摄影模式", "产品研发部 + 品牌部"],
        ["产品拍摄", "Day 1-4", "模式A（专职外拍）/ B（区域美食摄影师）/ C（门店实拍SOP），同步产出主图+详情图+短视频素材", "品牌部 + 外部摄影/视频团队"],
        ["元媒体素材制作", "Day 1-4（与拍摄并行）", "设计稿、海报、小红书/抖音/微信推文排版，以品牌VI模板批量生产", "社媒编辑 + 视觉设计"],
        ["培训材料制作", "Day 2-6（与素材并行）", "制作门店培训PPT/视频、SOP操作手册，上传绿茶大学平台", "培训部（Training Dept）"],
        ["门店培训执行", "Day 5-10", "直播培训（绿茶大学）/ 区域经理线下示范，确认门店master覆盖100%", "区域经理 + 绿茶大学平台"],
        ["私域预热", "Day 8-14", "微信群'周五福利日'预告，WOW会员定向push，KOL样品体验", "私域运营团队（数字化部）"],
        ["全渠道正式上线", "Day 15-20", "堂食菜单更新 + 外卖平台上线 + 公众号/小红书/抖音推文发布", "运营 + 数字化 + 品牌"],
    ],
    [4.0, 3.0, 6.0, 3.5]
)

subheading("产品拍摄三种模式")
add_table(
    ["模式", "适用情景", "时效", "质量"],
    [
        ["A · 专职外部摄影师", "大促重点菜、品牌形象菜", "3-5天", "高（商业大片级）"],
        ["B · 区域美食摄影合作伙伴", "中等优先级新菜，区域差异化菜品", "1-3天", "中（达人/KOL素材级）"],
        ["C · 门店实拍SOP", "常规迭代菜、限时特供菜，追求极致时效", "当日完成", "基础（符合平台发布标准）"],
    ],
    [4.5, 5.5, 2.5, 2.5]
)

subheading("社交渠道压缩机制")
add_table(
    ["渠道", "内容形式", "发布节点", "并行条件"],
    [
        ["抖音/视频号", "15-60秒短视频", "Day 12-15 上线前", "素材在Day4前完成剪辑"],
        ["小红书", "图文种草笔记", "Day 10-14 预热期", "摄影素材Day4前到位"],
        ["公众号", "新品推文（图文）", "Day 15-20 正式上线日", "以上素材复用"],
        ["私域微信群", "福利日活动通知 + 券", "Day 8-12 预热", "优惠信息在Day6前确认"],
        ["外卖平台（美团/饿了么）", "菜品图文上架", "Day 15-20 与堂食同步", "主图在Day5前完成"],
    ],
    [3.5, 4.0, 3.5, 4.5]
)

hr()

# ══════════════════════════════════════════════════════════════════════════════
# Q3: 营销活动上线前培训
# ══════════════════════════════════════════════════════════════════════════════
heading("Q3 · 营销活动上线前培训的执行体系与部门协作", 1)

confidence_badge("高置信 · HIGH CONFIDENCE", "high")
body("绿茶集团'绿茶大学'培训平台经公开媒体报道证实，直播培训形式有资料佐证。三级培训结构符合其三级组织架构（集团→区域→门店）的内在逻辑，置信度相对最高。")
doc.add_paragraph()

subheading("三级培训传导结构")
add_table(
    ["层级", "责任方", "对象", "形式"],
    [
        ["L1 · 集团级内容制作", "营销/品牌部 + 培训部（Training & HR）", "全体区域经理 + 门店经理", "绿茶大学平台直播 / 录播视频 / 文档SOP"],
        ["L2 · 区域级传导", "大区经理 / 区域督导", "辖区内所有门店经理 + 关键岗位（前厅/后厨负责人）", "区域集中培训 / 线上直播跟进 / 抽检考核"],
        ["L3 · 门店级落地", "门店店长（经理）", "全体门店员工（服务员、收银、厨房）", "门店晨会/夕会、实操演练、以老带新"],
    ],
    [3.5, 4.0, 5.0, 4.0]
)

subheading("五种培训形式")
add_table(
    ["培训形式", "适用场景", "时效性", "覆盖半径"],
    [
        ["绿茶大学直播培训", "全国性大促前、新品系统上线前", "T-7天前完成", "全国465+门店同步"],
        ["录播视频+考核", "常规迭代、非紧急但须合规的操作变更", "可随时回看", "自学型，须后台跟踪完成率"],
        ["区域集中培训（线下）", "大促、重大品牌活动、高复杂度新菜", "T-5天前完成", "覆盖区域内门店，确保经理级以上100%"],
        ["SOP文档/图解手册", "菜品出品标准、摆盘规范、话术脚本", "随物料包同步下发", "门店自学+抽查"],
        ["门店晨会/夕会传导", "最后一公里，确保一线员工知晓", "T-1天", "100%覆盖，责任到店长"],
    ],
    [4.5, 4.5, 3.0, 3.5]
)

subheading("培训视频制作责任方")
add_table(
    ["视频内容类型", "主要制作责任方", "参与部门", "产出形式"],
    [
        ["品牌主题讲解 / 活动背景介绍", "品牌部主导", "社媒/设计配合", "10-15min讲解视频"],
        ["新品菜品操作标准 / 出品规范", "培训部 + 产品研发部联合", "厨房技术团队演示", "3-8min操作演示视频"],
        ["门店执行SOP / 服务流程", "运营部 + 培训部联合", "督导示范", "5-10min场景演示视频"],
        ["私域话术 / 会员转化脚本", "数字化/私域部主导", "品牌部审核", "文字脚本 + 1-2min示范视频"],
        ["营销工具使用（POS/小程序）", "IT/数字化部主导", "运营部配合", "屏幕录制教程"],
    ],
    [4.5, 4.0, 3.5, 3.5]
)

callout("💡 绿茶大学平台作为数字化培训中枢，承接L1→L2传导的主要媒介，直播互动确保区域经理实时对齐；录播为L2→L3传导留存可回溯资产，是质量控制的关键机制。")

hr()

# ══════════════════════════════════════════════════════════════════════════════
# RISK SIGNALS
# ══════════════════════════════════════════════════════════════════════════════
heading("低置信风险信号 · 评审裁定", 1)

body("Critic 智能体识别的待验证项，需访谈中直接确认。")
doc.add_paragraph()

add_table(
    ["信号", "低置信原因", "建议验证问题"],
    [
        ["L4董事会审批触发条件", "金额门槛和品牌转向定义未见公开披露，上市时间短（2025年5月），治理机制仍在建立中", "请问品牌营销预算中，哪一级别的审批需要上董事会？是否有金额门槛或内容类型门槛？"],
        ["15-20天数字来源", "在财报、招股书、媒体报道中未找到明确的15-20天官方表述，是行业基准推断", "请问新品从内容制作到全渠道上线，内部有没有目标的执行时间窗口？大概是多少天？"],
        ["审批会议具体名称", "'品牌日历审定会'是推断命名，实际内部会议类型和固定频次未知", "品牌年度日历是否有固定的审批会议？会议叫什么？多久开一次？"],
    ],
    [4.5, 5.5, 5.5]
)

hr()

# ══════════════════════════════════════════════════════════════════════════════
# VERDICT
# ══════════════════════════════════════════════════════════════════════════════
heading("评审裁定", 1)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
p.paragraph_format.left_indent  = Cm(0.5)
pPr = p._element.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"), "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"), "0A0A0B")
pPr.append(shd)
run = p.add_run("绿茶集团已建立覆盖\"品牌策划→产品上新→门店培训\"的完整数字化执行闭环。203道新品/年的节奏印证了并行执行基础设施的存在；绿茶大学平台是培训体系的核心杠杆。三条延伸问题的框架性答案置信度中等，具体审批会议名称、15-20天官方窗口定义、L4触发门槛三项需一手验证。访谈建议优先询问王勤松/于丽影对产品营销协同机制的直接表述。")
set_font(run, name_zh="宋体", size=10.5, color=(241,239,234))

doc.add_paragraph()

# ── Footer ────────────────────────────────────────────────────────────────────
hr()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("本报告由 Merlin 咨询情报系统生成 · 仅供内部参考 · 请勿对外传播")
set_font(run, size=8, color=(150,150,150))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("反馈：https://spontaneous-youtiao-b9cde9.netlify.app")
set_font(run, size=8, color=(139,108,66))

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/home/user/work/reports/greenteagroup_Followup_Report.docx"
doc.save(out)
print(f"Saved: {out}")
